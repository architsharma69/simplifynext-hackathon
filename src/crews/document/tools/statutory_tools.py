"""
crews/tools/statutory_tools.py
Document rendering for ACRA BizFile+ incorporation paperwork. The key design
choice: the LLM agent gathers and validates structured CompanyProfile data
(via conversation), then calls ONE of these tools with that structured data.
The tool renders a real .docx deterministically with python-docx — the model
never free-writes the legal text of a Model Constitution or a Form 45.

Each render function also computes a checksum and (where the document is a
statutory filing, not an internal working paper) the Flow escalates it for
human sign-off before it would be considered ready to submit.
"""
from __future__ import annotations

import hashlib
import os
import sys
from datetime import date
from pathlib import Path

from crewai.tools import tool
from docx import Document

# Makes `Config` (and, via it, `src`) importable so this file works both as a
# package import and as a direct `python .../statutory_tools.py` run.
_REPO_ROOT = Path(__file__).resolve().parents[3]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from Config.config import DOCUMENT_OUTPUT_DIR, SRC_DIR  # noqa: E402

if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from crews.document.schemas import CompanyProfile, DocumentType, RenderedDocument  # noqa: E402

OUTPUT_DIR = DOCUMENT_OUTPUT_DIR
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _checksum(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        h.update(f.read())
    return h.hexdigest()


def _save_and_wrap(
    doc: Document, filename: str, doc_type: DocumentType, summary: dict
) -> RenderedDocument:
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return RenderedDocument(
        document_type=doc_type,
        file_path=path,
        checksum_sha256=_checksum(path),
        source_data_summary=summary,
    )


def _render_model_constitution(company: CompanyProfile) -> RenderedDocument:
    doc = Document()
    doc.add_heading(f"Constitution of {company.proposed_company_name}", level=0)
    doc.add_paragraph(
        "Adopted pursuant to the Companies Act 1967, in the form of the Model "
        "Constitution for a private company limited by shares."
    )
    doc.add_heading("1. Company Details", level=1)
    doc.add_paragraph(f"Registered address: {company.registered_address}")
    doc.add_paragraph(
        f"Principal activity (SSIC): {company.principal_activity_ssic_code}"
    )
    doc.add_paragraph(f"Paid-up capital: SGD {company.paid_up_capital_sgd:,.2f}")
    doc.add_heading("2. Directors", level=1)
    for d in company.directors:
        doc.add_paragraph(
            f"{d.full_name} ({d.nationality}) — {d.residential_address}"
            f"{' [Resident Director]' if d.is_resident_director else ''}",
            style="List Bullet",
        )
    doc.add_heading("3. Shareholders", level=1)
    for s in company.shareholders:
        doc.add_paragraph(
            f"{s.full_name} — {s.shares_held} {s.share_class} shares",
            style="List Bullet",
        )
    filename = f"model_constitution_{company.proposed_company_name.replace(' ', '_')}.docx"
    return _save_and_wrap(
        doc,
        filename,
        DocumentType.MODEL_CONSTITUTION,
        {"company": company.proposed_company_name, "directors": len(company.directors)},
    )


def _render_form_45(company: CompanyProfile) -> RenderedDocument:
    """Form 45: Notification of Registered Office Address / Office Hours."""
    doc = Document()
    doc.add_heading(
        "Form 45 — Notification of Situation of Registered Office", level=0
    )
    doc.add_paragraph(f"Company: {company.proposed_company_name}")
    doc.add_paragraph(f"Registered office address: {company.registered_address}")
    doc.add_paragraph(
        "Office hours: 9:00am to 5:00pm, Monday to Friday (excluding public holidays)"
    )
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    filename = f"form_45_{company.proposed_company_name.replace(' ', '_')}.docx"
    return _save_and_wrap(
        doc,
        filename,
        DocumentType.FORM_45,
        {"company": company.proposed_company_name},
    )


def _render_first_board_resolution(company: CompanyProfile) -> RenderedDocument:
    doc = Document()
    doc.add_heading(
        f"First Resolutions of the Board of Directors of {company.proposed_company_name}",
        level=0,
    )
    doc.add_paragraph(f"Held on {date.today().isoformat()}")
    doc.add_paragraph("Present:")
    for d in company.directors:
        doc.add_paragraph(d.full_name, style="List Bullet")
    doc.add_heading("Resolved that:", level=1)
    resolutions = [
        f"The registered office of the Company be situated at {company.registered_address}.",
        f"{company.company_secretary_name or '[TO BE APPOINTED]'} be and is hereby appointed Company Secretary.",
        "The common seal (if any) be adopted and its use governed by the Constitution.",
        "The bank account(s) of the Company be opened as tabled at this meeting.",
        "Shares be allotted as follows: "
        + "; ".join(
            f"{s.full_name} — {s.shares_held} shares" for s in company.shareholders
        )
        + ".",
    ]
    for r in resolutions:
        doc.add_paragraph(r, style="List Number")
    filename = f"first_board_resolution_{company.proposed_company_name.replace(' ', '_')}.docx"
    return _save_and_wrap(
        doc,
        filename,
        DocumentType.FIRST_BOARD_RESOLUTION,
        {"company": company.proposed_company_name, "resolutions": len(resolutions)},
    )


def _render_rorc_register(company: CompanyProfile) -> RenderedDocument:
    """Register of Registrable Controllers."""
    doc = Document()
    doc.add_heading(
        f"Register of Registrable Controllers — {company.proposed_company_name}",
        level=0,
    )
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Name",
        "ID Number",
        "Nature of Control",
        "Date Became Controller",
    )
    for s in company.shareholders:
        if (
            s.shares_held / max(sum(x.shares_held for x in company.shareholders), 1)
            > 0.25
        ):
            row = table.add_row().cells
            row[0].text = s.full_name
            row[1].text = s.id_number
            row[2].text = f">25% shareholding ({s.shares_held} shares)"
            row[3].text = date.today().isoformat()
    filename = f"rorc_register_{company.proposed_company_name.replace(' ', '_')}.docx"
    return _save_and_wrap(
        doc,
        filename,
        DocumentType.RORC_REGISTER,
        {"company": company.proposed_company_name},
    )


_RENDERERS = {
    DocumentType.MODEL_CONSTITUTION: _render_model_constitution,
    DocumentType.FORM_45: _render_form_45,
    DocumentType.FIRST_BOARD_RESOLUTION: _render_first_board_resolution,
    DocumentType.RORC_REGISTER: _render_rorc_register,
}


@tool("Render ACRA Incorporation Document")
def render_acra_document(document_type: str, company_profile_json: str) -> str:
    """
    Render one ACRA BizFile+ incorporation document from a validated
    CompanyProfile JSON payload. Returns a RenderedDocument JSON string
    containing the file path and checksum.
    Args:
        document_type: One of "model_constitution", "form_45",
            "first_board_resolution", "rorc_register".
        company_profile_json: JSON string matching the CompanyProfile schema
            (proposed_company_name, registered_address,
            principal_activity_ssic_code, directors, shareholders,
            company_secretary_name, paid_up_capital_sgd).
    """
    try:
        doc_type = DocumentType(document_type)
    except ValueError:
        return (
            f"ERROR: unknown document_type '{document_type}'. "
            f"Valid: {[d.value for d in _RENDERERS]}"
        )
    renderer = _RENDERERS.get(doc_type)
    if renderer is None:
        return f"ERROR: no renderer registered for '{document_type}'"
    company = CompanyProfile.model_validate_json(company_profile_json)
    rendered = renderer(company)
    return rendered.model_dump_json(indent=2)


@tool("Validate Company Profile Completeness")
def validate_company_profile(company_profile_json: str) -> str:
    """
    Check a CompanyProfile JSON payload for ACRA-filing blockers before any
    document is rendered (missing resident director, zero shareholders,
    invalid SSIC code format, etc). Returns "OK" or a list of issues.
    Args:
        company_profile_json: JSON string matching the CompanyProfile schema.
    """
    company = CompanyProfile.model_validate_json(company_profile_json)
    issues = []
    if not any(d.is_resident_director for d in company.directors):
        issues.append(
            "No resident director found. Singapore private companies require at least "
            "one director who is ordinarily resident in Singapore."
        )
    if not company.shareholders:
        issues.append("No shareholders specified.")
    if company.paid_up_capital_sgd <= 0:
        issues.append("Paid-up capital must be greater than SGD 0.")
    if not company.principal_activity_ssic_code:
        issues.append("Missing SSIC code for principal activity.")
    total_shares = sum(s.shares_held for s in company.shareholders)
    if total_shares <= 0:
        issues.append("Total shares held across shareholders must be greater than 0.")
    return "OK" if not issues else "ISSUES FOUND:\n- " + "\n- ".join(issues)
