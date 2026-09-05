"""
crews/tools/grant_tools.py
Compiles a structured grant package (Startup SG Founder / EDG) from the
outputs of the other two sub-agents (FinancialForecast, HeadcountPlan) plus
narrative sections the agent has drafted with the user. Rendering is
deterministic docx generation, same pattern as statutory_tools.
"""
from __future__ import annotations

import hashlib
import json
import os
import sys
from pathlib import Path

from crewai.tools import tool
from docx import Document

# Makes `Config` (and, via it, `src`) importable so this file works both as a
# package import and as a direct `python .../grant_tools.py` run.
_REPO_ROOT = Path(__file__).resolve().parents[3]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from Config.config import DOCUMENT_OUTPUT_DIR, SRC_DIR  # noqa: E402

if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from crews.document.schemas import (  # noqa: E402
    CompanyProfile,
    DocumentType,
    FinancialForecast,
    GrantPackage,
    GrantScheme,
    HeadcountPlan,
    RenderedDocument,
)

OUTPUT_DIR = DOCUMENT_OUTPUT_DIR
os.makedirs(OUTPUT_DIR, exist_ok=True)

_REQUIRED_SECTIONS = {
    GrantScheme.STARTUP_SG_FOUNDER: [
        "problem_statement",
        "solution",
        "market_opportunity",
        "founder_background",
        "use_of_funds",
    ],
    GrantScheme.EDG: [
        "project_scope",
        "capability_area",  # Core Capabilities / Innovation & Productivity / Market Access
        "expected_outcomes",
        "project_timeline",
        "use_of_funds",
    ],
}


def _checksum(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        h.update(f.read())
    return h.hexdigest()


@tool("Validate Grant Package Narrative Completeness")
def validate_grant_narrative(scheme: str, narrative_sections_json: str) -> str:
    """
    Check that all required narrative sections for a given grant scheme are
    present and non-trivial before compiling the final package.
    Args:
        scheme: "startup_sg_founder" or "enterprise_development_grant".
        narrative_sections_json: JSON object mapping section name -> text,
            e.g. {"problem_statement": "...", "solution": "..."}.
    """
    try:
        scheme_enum = GrantScheme(scheme)
    except ValueError:
        return f"ERROR: unknown scheme '{scheme}'. Valid: {[s.value for s in GrantScheme]}"
    sections = json.loads(narrative_sections_json)
    required = _REQUIRED_SECTIONS[scheme_enum]
    missing = [r for r in required if not sections.get(r, "").strip()]
    thin = [r for r in required if r not in missing and len(sections.get(r, "")) < 80]
    if not missing and not thin:
        return "OK"
    msg = []
    if missing:
        msg.append(f"Missing sections: {missing}")
    if thin:
        msg.append(f"Sections likely too thin (<80 chars) to be reviewable: {thin}")
    return "ISSUES FOUND:\n" + "\n".join(msg)


@tool("Compile Grant Package Document")
def compile_grant_package(
    scheme: str,
    company_profile_json: str,
    financial_forecast_json: str,
    headcount_plan_json: str,
    narrative_sections_json: str,
    requested_amount_sgd: float,
) -> str:
    """
    Assemble the final grant package .docx from validated component data.
    Should only be called after validate_grant_narrative returns "OK" and
    after the financial forecast / headcount plan already exist. Returns a
    GrantPackage JSON string with generated_document_path set.
    Args:
        scheme: "startup_sg_founder" or "enterprise_development_grant".
        company_profile_json: JSON string matching CompanyProfile.
        financial_forecast_json: JSON string matching FinancialForecast
            (from generate_financial_forecast).
        headcount_plan_json: JSON string matching HeadcountPlan.
        narrative_sections_json: JSON object of section name -> text.
        requested_amount_sgd: Grant amount being requested, in SGD.
    """
    scheme_enum = GrantScheme(scheme)
    company = CompanyProfile.model_validate_json(company_profile_json)
    financials = FinancialForecast.model_validate_json(financial_forecast_json)
    headcount = HeadcountPlan.model_validate_json(headcount_plan_json)
    sections: dict[str, str] = json.loads(narrative_sections_json)

    doc = Document()
    title = (
        "Startup SG Founder Application"
        if scheme_enum == GrantScheme.STARTUP_SG_FOUNDER
        else "Enterprise Development Grant (EDG) Application"
    )
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Applicant: {company.proposed_company_name}")
    doc.add_paragraph(f"Requested amount: SGD {requested_amount_sgd:,.2f}")
    for section_name in _REQUIRED_SECTIONS[scheme_enum]:
        doc.add_heading(section_name.replace("_", " ").title(), level=1)
        doc.add_paragraph(sections.get(section_name, "[NOT PROVIDED]"))

    doc.add_heading("Financial Summary", level=1)
    doc.add_paragraph(
        f"Monthly burn rate: SGD {financials.monthly_burn_rate_sgd:,.2f}. "
        f"Break-even month: {financials.break_even_month_index or 'not within 36 months'}. "
        f"3-year cumulative revenue: SGD {financials.three_year_revenue_total:,.2f}."
    )

    doc.add_heading("Headcount Plan", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Role",
        "Dept",
        "Count",
        "Monthly Salary (SGD)",
    )
    for line in headcount.lines:
        row = table.add_row().cells
        row[0].text = line.role_title
        row[1].text = line.department
        row[2].text = str(line.count)
        row[3].text = f"{line.monthly_salary_sgd:,.2f}"

    filename = (
        f"grant_{scheme_enum.value}_{company.proposed_company_name.replace(' ', '_')}.docx"
    )
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)

    rendered = RenderedDocument(
        document_type=(
            DocumentType.GRANT_PACKAGE_STARTUP_SG
            if scheme_enum == GrantScheme.STARTUP_SG_FOUNDER
            else DocumentType.GRANT_PACKAGE_EDG
        ),
        file_path=path,
        checksum_sha256=_checksum(path),
        source_data_summary={
            "company": company.proposed_company_name,
            "scheme": scheme_enum.value,
        },
    )
    package = GrantPackage(
        scheme=scheme_enum,
        company=company,
        financials=financials,
        headcount=headcount,
        narrative_sections=sections,
        requested_amount_sgd=requested_amount_sgd,
        generated_document_path=rendered.file_path,
    )
    return package.model_dump_json(indent=2)
